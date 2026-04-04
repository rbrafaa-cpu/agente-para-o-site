"""
fetch_kb.py — Step 1 of the RAG pipeline

Exports the Google Doc as DOCX, extracts text chunks and all embedded images
(as base64 data URIs), and saves the result to .tmp/kb_chunks.json.

Images are pulled directly from the DOCX ZIP archive — no external HTTP auth needed.

Usage:
    python tools/fetch_kb.py
"""
from __future__ import annotations

import base64
import json
import re
import sys
from io import BytesIO
from pathlib import Path

import httpx
from docx import Document
from docx.oxml.ns import qn
from dotenv import load_dotenv
from PIL import Image

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
load_dotenv()

DOC_ID = "13Vrr76mp5RpQwYdnOFOStQHIyXqeiHJOotKLbOsWL4w"
DOCX_EXPORT_URL = f"https://docs.google.com/document/d/{DOC_ID}/export?format=docx"
OUTPUT_PATH = Path(__file__).parent.parent / ".tmp" / "kb_chunks.json"
MAX_CHUNK_CHARS = 1800  # ~500 tokens

HEADING_STYLES = {"heading 1", "heading 2", "heading 3", "title"}
MAX_IMAGE_PX = 320     # max width or height — chat thumbnails don't need full res
JPEG_QUALITY = 50      # keeps base64 well under Pinecone's 40KB metadata limit


# ---------------------------------------------------------------------------
# Fetch DOCX
# ---------------------------------------------------------------------------

LOCAL_DOCX_GLOB = "*.docx"  # any .docx in project root (excludes ~$ lock files)


def find_local_docx() -> Path | None:
    """Return the first non-lock .docx file in the project root, if any."""
    project_root = Path(__file__).parent.parent
    candidates = [p for p in project_root.glob("*.docx") if not p.name.startswith("~$")]
    return candidates[0] if candidates else None


def fetch_docx_bytes() -> bytes:
    """Use local .docx if present, otherwise download from Google."""
    local = find_local_docx()
    if local:
        print(f"Using local file: {local.name}")
        data = local.read_bytes()
        print(f"  {len(data):,} bytes")
        return data

    print(f"Fetching DOCX: {DOCX_EXPORT_URL}")
    try:
        r = httpx.get(DOCX_EXPORT_URL, follow_redirects=True, timeout=60)
        if r.status_code == 200 and "accounts.google.com" not in str(r.url):
            print(f"  Downloaded {len(r.content):,} bytes")
            return r.content
        print("  Auth redirect — falling back to GWS CLI...")
    except Exception as e:
        print(f"  HTTP error: {e} — falling back to GWS CLI...")

    return fetch_docx_via_gws()


def fetch_docx_via_gws() -> bytes:
    """Use the GWS CLI to export the doc as DOCX."""
    import subprocess, tempfile, os
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name

    result = subprocess.run(
        ["gws", "docs", "export", DOC_ID, "--format", "docx", "--output", tmp_path],
        capture_output=True, text=True, timeout=60,
    )
    if result.returncode != 0:
        print(f"GWS CLI error: {result.stderr}")
        sys.exit(1)

    data = Path(tmp_path).read_bytes()
    os.unlink(tmp_path)
    print(f"  GWS CLI downloaded {len(data):,} bytes")
    return data


# ---------------------------------------------------------------------------
# Image helpers
# ---------------------------------------------------------------------------

def compress_image(raw_bytes: bytes) -> str:
    """
    Resize + compress image to JPEG, return as base64 data URI.
    Keeps output well under Pinecone's 40KB metadata limit.
    """
    try:
        img = Image.open(BytesIO(raw_bytes)).convert("RGB")
        img.thumbnail((MAX_IMAGE_PX, MAX_IMAGE_PX), Image.LANCZOS)
        buf = BytesIO()
        img.save(buf, format="JPEG", quality=JPEG_QUALITY, optimize=True)
        b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
        return f"data:image/jpeg;base64,{b64}"
    except Exception as e:
        print(f"    Image compression failed: {e}")
        return ""


def extract_images_from_para(para, doc: Document) -> list[str]:
    """Return compressed base64 data URIs for any images embedded in this paragraph."""
    images = []
    for drawing in para._element.findall(".//" + qn("w:drawing")):
        for blip in drawing.findall(".//" + qn("a:blip")):
            rId = blip.get(qn("r:embed"))
            if rId and rId in doc.part.related_parts:
                img_part = doc.part.related_parts[rId]
                data_uri = compress_image(img_part.blob)
                if data_uri:
                    size_kb = len(data_uri) / 1024
                    print(f"    Captured image ({size_kb:.1f} KB compressed)")
                    images.append(data_uri)
    return images


# ---------------------------------------------------------------------------
# Text helpers
# ---------------------------------------------------------------------------

def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def split_into_chunks(text: str, max_chars: int = MAX_CHUNK_CHARS) -> list[str]:
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    chunks, current = [], ""
    for para in paragraphs:
        if len(current) + len(para) + 2 <= max_chars:
            current = (current + "\n\n" + para).strip()
        else:
            if current:
                chunks.append(current)
            if len(para) > max_chars:
                sentences = re.split(r"(?<=[.!?])\s+", para)
                sub = ""
                for sent in sentences:
                    if len(sub) + len(sent) + 1 <= max_chars:
                        sub = (sub + " " + sent).strip()
                    else:
                        if sub:
                            chunks.append(sub)
                        sub = sent
                if sub:
                    chunks.append(sub)
            else:
                current = para
    if current:
        chunks.append(current)
    return chunks or [""]


# ---------------------------------------------------------------------------
# Parse DOCX into chunks
# ---------------------------------------------------------------------------

def parse_docx(docx_bytes: bytes) -> list[dict]:
    doc = Document(BytesIO(docx_bytes))

    chunks: list[dict] = []
    chunk_index = 0
    current_section = "Introduction"
    current_text_parts: list[str] = []
    # Each entry: {"data": base64_data_uri, "caption": str}
    current_images: list[dict] = []
    total_images = 0

    def flush():
        nonlocal chunk_index
        text = clean_text("\n\n".join(current_text_parts))
        if not text and not current_images:
            return
        sub_chunks = split_into_chunks(text)
        for i, sub_text in enumerate(sub_chunks):
            chunks.append({
                "chunk_index": chunk_index,
                "section_title": current_section,
                "text": sub_text,
                "images": current_images[:] if i == 0 else [],
            })
            chunk_index += 1
        current_text_parts.clear()
        current_images.clear()

    # Use indexed iteration so we can look ahead for captions
    paragraphs = list(doc.paragraphs)
    for i, para in enumerate(paragraphs):
        style_name = (para.style.name or "").lower()
        text = clean_text(para.text)
        imgs = extract_images_from_para(para, doc)
        total_images += len(imgs)

        if style_name in HEADING_STYLES:
            flush()
            if text:
                current_section = text
        else:
            if text:
                current_text_parts.append(text)
            if imgs:
                # Look ahead: if next paragraph is a caption style or short text, use it
                caption = ""
                if i + 1 < len(paragraphs):
                    next_para = paragraphs[i + 1]
                    next_style = (next_para.style.name or "").lower()
                    next_text = clean_text(next_para.text)
                    if "caption" in next_style or (next_text and len(next_text) < 200 and not extract_images_from_para(next_para, doc)):
                        caption = next_text
                # If no caption found, fall back to current paragraph text or section title
                if not caption:
                    caption = text or current_section
                for img_b64 in imgs:
                    current_images.append({"data": img_b64, "caption": caption})

    flush()

    # Walk tables — associate each cell's text with its images as captions
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text_parts: list[str] = []
                cell_imgs: list[str] = []
                for cell_para in cell.paragraphs:
                    t = clean_text(cell_para.text)
                    if t:
                        cell_text_parts.append(t)
                    cell_imgs.extend(extract_images_from_para(cell_para, doc))
                cell_caption = " ".join(cell_text_parts)
                current_text_parts.extend(cell_text_parts)
                for img_b64 in cell_imgs:
                    current_images.append({"data": img_b64, "caption": cell_caption or current_section})
        flush()

    chunks = [c for c in chunks if c["text"] or c["images"]]
    print(f"  Produced {len(chunks)} chunks, {total_images} images captured")
    return chunks


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    docx_bytes = fetch_docx_bytes()
    chunks = parse_docx(docx_bytes)

    with open(OUTPUT_PATH, "w", encoding="utf-8") as f:
        json.dump(chunks, f, ensure_ascii=False, indent=2)

    print(f"Saved {len(chunks)} chunks → {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
